from docx import Document
from io import BytesIO


def create_improved_docx(original_text, improved_sections):
    doc = Document()
    doc.add_heading('Improved Curriculum Vitae', 0)

    doc.add_heading('AI Suggestions & Improvements', level=1)
    for section in improved_sections:
        p = doc.add_paragraph()
        p.add_run(f"Reason: ").bold = True
        p.add_run(section.get('explanation', ''))

        p = doc.add_paragraph()
        p.add_run(f"Original: ").italic = True
        p.add_run(section.get('original', ''))

        p = doc.add_paragraph()
        p.add_run(f"Improved Version: ").bold = True
        p.add_run(section.get('improved', ''))
        doc.add_paragraph("-" * 20)

    doc.add_page_break()
    doc.add_heading('Full Text Content', level=1)
    doc.add_paragraph(original_text)

    # שמירה לזיכרון (Buffer) במקום לקובץ פיזי
    target = BytesIO()
    doc.save(target)
    target.seek(0)
    return target