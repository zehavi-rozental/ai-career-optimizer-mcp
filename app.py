import streamlit as st
from services.google_search import GoogleSearchService
from services.ai_service import AIService
from utils.pdf_processor import extract_text_from_pdf, extract_text_from_docx
from docx import Document
from io import BytesIO
import re

st.set_page_config(page_title="AI Career Optimizer Pro", page_icon="🎯", layout="wide")

# אתחול Session State
if 'search_results' not in st.session_state: st.session_state.search_results = []
if 'search_query' not in st.session_state: st.session_state.search_query = ""
if 'resume_text' not in st.session_state: st.session_state.resume_text = ""
if 'analysis_results' not in st.session_state: st.session_state.analysis_results = None
# המפתח שישלוט בתיבת הטקסט
if 'main_job_area' not in st.session_state: st.session_state.main_job_area = ""

st.title("🎯 AI Career Optimizer Pro")

# --- שלב 1: העלאת קורות חיים ---
st.markdown("## 📄 שלב 1: העלאת קורות חיים")
uploaded_file = st.file_uploader("העלי קורות חיים (PDF או DOCX):", type=["pdf", "docx"])
if uploaded_file:
    with st.spinner("מעבד קורות חיים..."):
        if uploaded_file.name.endswith('.pdf'):
            st.session_state.resume_text = extract_text_from_pdf(uploaded_file)
        else:
            st.session_state.resume_text = extract_text_from_docx(uploaded_file)
    st.success("✅ קורות חיים נטענו בהצלחה!")

st.divider()

# --- שלב 2: חיפוש משרות ---
st.markdown("## 🔍 שלב 2: חיפוש משרות ספציפיות")
col_search, col_btn = st.columns([4, 1])
with col_search:
    job_query = st.text_input("איזו משרה את מחפשת?", value=st.session_state.search_query)
with col_btn:
    st.write(" ")
    if st.button("🔎 חפש משרות", use_container_width=True):
        st.session_state.search_query = job_query
        with st.spinner("מחפש..."):
            st.session_state.search_results = GoogleSearchService.search_jobs(job_query, st.session_state.resume_text)

if st.session_state.search_results:
    for i, job in enumerate(st.session_state.search_results):
        score = job.get('quick_score', 0)
        color = "#2ecc71" if score > 70 else "#f1c40f"

        st.markdown(f"""
        <div style="border:1px solid #ddd; padding:15px; border-radius:10px; margin-bottom:10px; background-color:#f9f9f9;">
            <b style="font-size:16px;">📌 {job.get('title')}</b> | <span style="color:{color}; font-weight:bold;">{score}% התאמה</span>
            <p style="font-size:14px; margin-top:5px;">{job.get('snippet', '')}</p>
        </div>
        """, unsafe_allow_html=True)

        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button(f"⬇️ שאוב תיאור מלא", key=f"fetch_{i}"):
                with st.spinner("שואב נתונים..."):
                    # שליפה ישירה של התוכן המלא
                    full_description = GoogleSearchService.get_full_job_content(job.get('link'))

                    # אם השאיבה הצליחה, נכניס אותה. אם לא, נשתמש בתקציר הקיים
                    final_text = full_description if len(full_description) > 150 else job.get('snippet', '')

                    st.session_state.main_job_area = final_text
                    st.rerun()
        with c2:
            if st.button(f"✅ בחר תקציר", key=f"sel_{i}"):
                # עדכון ישיר של ה-Key של ה-widget
                st.session_state.main_job_area = job.get('snippet', '')
                st.rerun()
        with c3:
            st.link_button("🔗 למשרה באתר", job.get('link', '#'))

st.divider()

# --- שלב 3: ניתוח התאמה ושיפור ---
st.markdown("## 📊 שלב 3: ניתוח התאמה ושיפור")

# יצירת התיבה שמחוברת ישירות ל-Session State
job_desc_input = st.text_area(
    "תיאור המשרה לניתוח:",
    height=300,
    key="main_job_area"  # קישור ישיר למפתח בזיכרון
)

if st.button("🚀 נתחי התאמה עכשיו", type="primary"):
    # שימוש בטקסט שיש כרגע בזיכרון התיבה
    current_desc = st.session_state.main_job_area

    if not st.session_state.resume_text:
        st.error("אנא העלי קורות חיים בשלב 1!")
    elif not current_desc or len(current_desc) < 10:
        st.error("המלבן ריק! אנא שאבי משרה או הדביקי תיאור ידנית.")
    else:
        with st.spinner("AI מנתח התאמה ומכין המלצות..."):
            analysis = AIService.analyze_job_match(st.session_state.resume_text, current_desc)
            st.session_state.analysis_results = analysis
            st.rerun()

if st.session_state.analysis_results:
    st.markdown("### 📋 תוצאות הניתוח")
    st.markdown(st.session_state.analysis_results, unsafe_allow_html=True)

    # ייצוא ל-Word
    doc = Document()
    doc.add_heading('ניתוח התאמה למשרה', 0)
    clean_text = re.sub('<[^<]+?>', '', st.session_state.analysis_results)
    doc.add_paragraph(clean_text)
    bio = BytesIO()
    doc.save(bio)
    st.download_button("📥 הורדי ניתוח כ-Word", bio.getvalue(), "analysis.docx")